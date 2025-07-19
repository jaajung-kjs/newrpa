import flet as ft
import pyautogui
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import datetime
import io

def create_base_report_document():
    document = Document()
    # 요청된 양식에 맞게 문서 생성
    # 1. 제목 추가: "공용노트북 보안점검 증빙사진" (글씨체 20, bold, 가운데 정렬)
    title_paragraph = document.add_paragraph()
    title_run = title_paragraph.add_run('공용노트북 보안점검 증빙사진')
    title_run.font.size = Pt(20)
    title_run.bold = True
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. 2x18 테이블 생성
    table = document.add_table(rows=18, cols=2)
    table.autofit = False # 자동 맞춤 해제
    # 표 열 너비 설정 (예시: 각 열을 3.2인치로 설정)
    table.columns[0].width = Inches(3.2)
    table.columns[1].width = Inches(3.2)

    # 테이블 스타일 설정 (기본 테두리 있는 스타일 사용)
    table.style = 'Table Grid'
    
    # 모든 셀의 내용을 가운데 정렬
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 3. 표 내용 채우기
    check_items = [
        "공유폴더 삭제 확인", "모델명 및 S/N", "OS 버전 및 설치일", "MAC 주소",
        "CMOS 패스워드", "윈도우 패스워드", "화면보호기", "업무자료 관리(인터넷 사용 PC)",
        "불필요 프로그램 삭제", "안티 바이러스(백신) 업데이트", "기타 프로그램 업데이트",
        "필수 보안 프로그램 설치 여부", "실시간 감시(V3)", "실시간 감시(Windows Defender)",
        "정밀검사(V3)", "정밀검사(Windows Defender)", "감염이력(V3)", "감염이력(Windows Defender)"
    ]
    
    current_item_index = 0
    for i in range(18): # 0부터 17까지 모든 행 반복
        if i % 2 == 0: # 짝수 인덱스 행 (0, 2, 4, ...)에 번호와 항목 채우기
            if current_item_index < len(check_items):
                table.cell(i, 0).paragraphs[0].text = f'{current_item_index + 1}. {check_items[current_item_index]}'
                current_item_index += 1
            else:
                table.cell(i, 0).paragraphs[0].text = ''

            if current_item_index < len(check_items):
                table.cell(i, 1).paragraphs[0].text = f'{current_item_index + 1}. {check_items[current_item_index]}'
                current_item_index += 1
            else:
                table.cell(i, 1).paragraphs[0].text = ''
        else: # 홀수 인덱스 행 (1, 3, 5, ...)은 사진을 위해 비워두기
            table.cell(i, 0).paragraphs[0].text = ''
            table.cell(i, 1).paragraphs[0].text = ''
    
    return document

def insert_image_to_document_table(document, image_object, row, col, width=Inches(3.0)):
    """
    주어진 Document 객체의 특정 테이블 셀에 이미지 객체를 삽입합니다.
    :param document: 이미지를 삽입할 docx.Document 객체
    :param image_object: 삽입할 PIL.Image.Image 객체
    :param row: 이미지를 삽입할 테이블 행 인덱스 (0-based)
    :param col: 이미지를 삽입할 테이블 열 인덱스 (0-based)
    :param width: 삽입할 이미지의 너비 (Inches 객체)
    """
    table = document.tables[0] # 첫 번째 테이블을 가정
    cell = table.cell(row, col)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    
    # 이미지 객체를 BytesIO에 저장하여 add_picture에 전달
    img_byte_arr = io.BytesIO()
    image_object.save(img_byte_arr, format='PNG') # PNG 형식으로 저장
    img_byte_arr.seek(0)
    run.add_picture(img_byte_arr, width=width)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER # 이미지 포함 단락 가운데 정렬


def run_rpa(page: ft.Page, status_text: ft.Text, start_button: ft.ElevatedButton):
    status_text.value = "점검중입니다. 아무것도 건드리지 마세요..."
    status_text.color = "red"
    start_button.disabled = True
    page.update() # UI 업데이트 강제

    try:
        # 스크린샷 캡처 로직은 main.py에서 처리

        # docx 문서 생성
        document = create_base_report_document()
        
        # 임시 스크린샷을 삽입하여 테스트 (main.py 통합 시 제거)
        # script_dir = os.path.dirname(os.path.abspath(__file__))
        # temp_screenshot_path = os.path.join(script_dir, "temp_screenshot.png")
        # pyautogui.screenshot(temp_screenshot_path)
        # insert_image_to_document_table(document, temp_screenshot_path, 1, 0) # 예시로 1행 0열에 삽입
        # os.remove(temp_screenshot_path) # 임시 파일 삭제

        # 3. docx 파일 저장
        output_filename = f"노트북_점검_보고서_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        document.save(output_filename)
        print(f"문서 저장: {output_filename}")

        status_text.value = f"점검 완료! 보고서가 '{output_filename}'(으)로 저장되었습니다."
        status_text.color = "green"
        page.update()

        def close_dialog(e):
            page.dialog.open = False
            page.update()

        page.dialog = ft.AlertDialog(title=ft.Text("완료"), content=ft.Text(f"노트북 점검이 완료되었습니다.\n보고서: {output_filename}"), actions=[ft.TextButton("확인", on_click=close_dialog)])
        page.dialog.open = True
        page.update()

    except Exception as e:
        status_text.value = f"오류 발생: {e}"
        status_text.color = "red"
        page.update()

        def close_dialog(e):
            page.dialog.open = False
            page.update()

        page.dialog = ft.AlertDialog(title=ft.Text("오류"), content=ft.Text(f"점검 중 오류가 발생했습니다: {e}"), actions=[ft.TextButton("확인", on_click=close_dialog)])
        page.dialog.open = True
        page.update()
    finally:
        start_button.disabled = False
        page.update()

async def main(page: ft.Page):
    page.title = "노트북 일괄 점검 RPA"
    page.vertical_alignment = ft.MainAxisAlignment.CENTER
    page.horizontal_alignment = ft.CrossAxisAlignment.CENTER
    page.window_width = 400
    page.window_height = 200
    page.window_resizable = False

    title_text = ft.Text("노트북 점검 프로그램", size=20, weight=ft.FontWeight.BOLD)
    status_text = ft.Text("시작하려면 버튼을 클릭하세요.", size=12, color="blue")
    start_button = ft.ElevatedButton("노트북 점검 DOCX 생성", on_click=lambda e: page.run_thread(run_rpa, page, status_text, start_button))

    page.add(
        title_text,
        status_text,
        start_button,
    )

if __name__ == "__main__":
    ft.app(target=main)
